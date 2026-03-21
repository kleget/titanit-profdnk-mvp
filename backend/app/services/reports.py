from __future__ import annotations

from dataclasses import dataclass
from datetime import timezone
from io import BytesIO
from typing import Literal

from docx import Document

from app.models import QuestionType, Submission, Test
from app.services.client_fields import REPORT_BLOCK_LABELS, normalize_client_fields_config
from app.web import templates

ReportKind = Literal["client", "psychologist"]


@dataclass
class AnswerRow:
    section_title: str
    question_text: str
    answer_value: str
    question_type: str


def _format_answer(value: object, question_type: QuestionType) -> str:
    if value is None:
        return "-"
    if question_type == QuestionType.YES_NO:
        return "Да" if value in {True, "true", "True", "1", 1, "yes", "on"} else "Нет"
    if isinstance(value, list):
        if not value:
            return "-"
        return ", ".join(str(v) for v in value)
    return str(value)


def _safe_float(value: object) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _build_chart_items(metrics: dict, derived_metrics: list[dict]) -> list[dict]:
    chart_items: list[dict] = []

    completion_percent = _safe_float(metrics.get("completion_percent")) or 0.0
    score_percent = _safe_float(metrics.get("score_percent")) or 0.0
    chart_items.append(
        {
            "label": "Заполнение теста",
            "value": round(completion_percent, 2),
            "percent": max(0.0, min(100.0, completion_percent)),
        }
    )
    chart_items.append(
        {
            "label": "Процент от максимума",
            "value": round(score_percent, 2),
            "percent": max(0.0, min(100.0, score_percent)),
        }
    )

    ok_derived = [
        metric for metric in derived_metrics if metric.get("status") == "ok" and metric.get("value") is not None
    ]
    if ok_derived:
        max_abs = max(abs(float(metric["value"])) for metric in ok_derived)
        scale = max(1.0, max_abs)
        for metric in ok_derived:
            numeric = float(metric["value"])
            chart_items.append(
                {
                    "label": metric.get("label") or metric.get("key"),
                    "value": round(numeric, 2),
                    "percent": max(0.0, min(100.0, (abs(numeric) / scale) * 100)),
                }
            )
    return chart_items


def _build_client_profile_rows(test: Test, submission: Submission) -> list[dict[str, str]]:
    config = normalize_client_fields_config(test.required_client_fields)
    required_builtin = set(config["required_builtin_fields"])

    extra_raw = submission.client_extra_json if isinstance(submission.client_extra_json, dict) else {}
    age_value = extra_raw.get("age")
    custom_values_raw = extra_raw.get("custom_fields")
    custom_values = custom_values_raw if isinstance(custom_values_raw, dict) else {}

    rows: list[dict[str, str]] = [
        {"label": "ФИО", "value": submission.client_full_name or "-"},
        {"label": "Email", "value": submission.client_email or "-"},
        {"label": "Телефон", "value": submission.client_phone or "-"},
    ]
    if age_value not in {None, ""} or "age" in required_builtin:
        rows.append({"label": "Возраст", "value": str(age_value or "-")})

    for field in config["custom_fields"]:
        key = str(field.get("key", "")).strip()
        label = str(field.get("label", "")).strip()
        if not key or not label:
            continue
        value = custom_values.get(key)
        if value in {None, ""} and not field.get("required"):
            continue
        rows.append({"label": label, "value": str(value or "-")})
    return rows


def _resolve_report_blocks(test: Test, report_kind: ReportKind) -> list[dict[str, str]]:
    config = normalize_client_fields_config(test.required_client_fields)
    templates_config = config.get("report_templates")
    if isinstance(templates_config, dict):
        block_keys_raw = templates_config.get(report_kind)
    else:
        block_keys_raw = None
    block_keys = block_keys_raw if isinstance(block_keys_raw, list) else []
    result: list[dict[str, str]] = []
    for key in block_keys:
        block_key = str(key).strip()
        label = REPORT_BLOCK_LABELS.get(block_key)
        if not label:
            continue
        result.append({"key": block_key, "label": label})
    return result


def build_report_context(
    test: Test,
    submission: Submission,
    report_kind: ReportKind,
) -> dict:
    answer_by_question_id = {answer.question_id: answer.value_json for answer in submission.answers}
    rows: list[AnswerRow] = []
    for section in test.sections:
        for question in section.questions:
            rows.append(
                AnswerRow(
                    section_title=section.title,
                    question_text=question.text,
                    answer_value=_format_answer(
                        answer_by_question_id.get(question.id), question.question_type
                    ),
                    question_type=question.question_type.value,
                )
            )

    metrics = submission.metrics_json or {}
    derived_metrics = metrics.get("derived_metrics") or []
    title = (
        "Клиентский отчёт"
        if report_kind == "client"
        else "Отчёт для профориентолога"
    )
    submitted_local = submission.submitted_at.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    client_profile_rows = _build_client_profile_rows(test, submission)
    report_blocks = _resolve_report_blocks(test, report_kind)
    return {
        "title": title,
        "report_kind": report_kind,
        "test_title": test.title,
        "test_description": test.description,
        "psychologist_name": test.psychologist.full_name,
        "client_name": submission.client_full_name,
        "client_email": submission.client_email or "-",
        "client_phone": submission.client_phone or "-",
        "client_profile_rows": client_profile_rows,
        "submitted_at": submitted_local,
        "metrics": metrics,
        "derived_metrics": derived_metrics,
        "chart_items": _build_chart_items(metrics, derived_metrics),
        "answers": rows,
        "report_blocks": report_blocks,
        "report_block_keys": [item["key"] for item in report_blocks],
    }


def render_html_report(context: dict, report_kind: ReportKind) -> str:
    template_name = (
        "reports/client_report.html"
        if report_kind == "client"
        else "reports/psychologist_report.html"
    )
    template = templates.env.get_template(template_name)
    return template.render(**context)


def build_docx_report(context: dict, report_kind: ReportKind) -> BytesIO:
    def append_profile_block() -> None:
        heading = "Данные анкеты" if report_kind == "client" else "Данные клиента"
        doc.add_heading(heading, level=1)
        for row in context.get("client_profile_rows", []):
            doc.add_paragraph(f"{row.get('label')}: {row.get('value')}")

    def append_summary_metrics_block() -> None:
        metrics = context.get("metrics", {})
        doc.add_heading("Метрики", level=1)
        doc.add_paragraph(
            f"Итоговый балл: {metrics.get('total_score', 0)} / {metrics.get('max_score', 0)}"
        )
        doc.add_paragraph(f"Заполнено: {metrics.get('completion_percent', 0)}%")
        doc.add_paragraph(f"Процент от максимума: {metrics.get('score_percent', 0)}%")

    def append_charts_block() -> None:
        chart_items = context.get("chart_items", [])
        doc.add_heading("Визуализация метрик", level=1)
        if not chart_items:
            doc.add_paragraph("Нет данных для визуализации.")
            return
        for item in chart_items:
            doc.add_paragraph(f"{item.get('label')}: {item.get('value')}")

    def append_derived_metrics_block() -> None:
        derived_metrics = context.get("derived_metrics", [])
        doc.add_heading("Производные метрики", level=1)
        if not derived_metrics:
            doc.add_paragraph("Производные метрики не настроены.")
            return
        for metric in derived_metrics:
            label = metric.get("label") or metric.get("key")
            value = metric.get("value")
            if metric.get("status") == "error":
                doc.add_paragraph(f"{label}: ошибка расчёта ({metric.get('error', 'unknown')})")
                continue
            doc.add_paragraph(f"{label}: {value}")

    def append_answers_block() -> None:
        doc.add_heading("Ответы", level=1)
        for row in context["answers"]:
            doc.add_paragraph(f"[{row.section_title}] {row.question_text}", style="List Bullet")
            doc.add_paragraph(row.answer_value)

    docx_block_renderers = {
        "profile": append_profile_block,
        "summary_metrics": append_summary_metrics_block,
        "charts": append_charts_block,
        "derived_metrics": append_derived_metrics_block,
        "answers": append_answers_block,
    }

    doc = Document()
    doc.add_heading(context["title"], level=0)
    doc.add_paragraph(f"Тест: {context['test_title']}")
    doc.add_paragraph(f"Психолог: {context['psychologist_name']}")
    doc.add_paragraph(f"Клиент: {context['client_name']}")
    doc.add_paragraph(f"Дата прохождения: {context['submitted_at']}")
    for block_key in context.get("report_block_keys", []):
        render = docx_block_renderers.get(str(block_key))
        if render:
            render()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
